
import re
import os
import tkinter as tk
from tkinter import filedialog, scrolledtext
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml

# --- Core Conversion Logic (from previous steps) ---

def parse_ass_file(ass_file_path):
    """
    Parses a .ass file to extract dialogue lines, including the speaker's name.
    """
    dialogue_lines = []
    dialogue_pattern = re.compile(r"Dialogue: [^,]+,[^,]+,[^,]+,[^,]+,([^,]*),[^,]+,[^,]+,[^,]+,[^,]*,?(.*)$")
    try:
        with open(ass_file_path, "r", encoding="utf-8") as f:
            for line in f:
                match = dialogue_pattern.match(line)
                if match:
                    name = match.group(1).strip()
                    dialogue_text = match.group(2).strip()
                    dialogue_text = re.sub(r"{[^}]+}", "", dialogue_text)
                    dialogue_text = dialogue_text.replace("\\N", " ")
                    dialogue_lines.append((name, dialogue_text))
    except Exception:
        # Return what was parsed, even if an error occurs
        return dialogue_lines
    return dialogue_lines

def create_docx(dialogue_lines, docx_file_path):
    """
    Creates a .docx file and writes the dialogue lines to it,
    grouping consecutive lines from the same speaker and aligning text.
    """
    if not dialogue_lines:
        return False

    document = Document()
    document.add_heading("Subtitle Transcript", 0)

    grouped_dialogues = []
    if dialogue_lines:
        current_speaker, current_text = dialogue_lines[0]
        for next_speaker, next_text in dialogue_lines[1:]:
            if next_speaker == current_speaker:
                current_text += f"\n{next_text}"
            else:
                grouped_dialogues.append((current_speaker, current_text))
                current_speaker = next_speaker
                current_text = next_text
        grouped_dialogues.append((current_speaker, current_text))

    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    tbl_props = table._tbl.tblPr
    tbl_borders = parse_xml(r'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/></w:tblBorders>')
    tbl_props.append(tbl_borders)

    table.columns[0].width = Pt(100)
    table.columns[1].width = Pt(400)

    for name, text in grouped_dialogues:
        row_cells = table.add_row().cells
        name_paragraph = row_cells[0].paragraphs[0]
        run = name_paragraph.add_run(f"{name}:")
        run.bold = True
        row_cells[1].text = text

    try:
        document.save(docx_file_path)
        return True
    except Exception:
        return False

# --- GUI Application ---

class SubtitleConverterApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Subtitle to DOCX Converter")
        self.root.geometry("600x400")

        self.input_files = []
        self.output_dir = ""

        # Frame for buttons
        button_frame = tk.Frame(root)
        button_frame.pack(pady=10)

        # Buttons
        self.select_files_btn = tk.Button(button_frame, text="1. Select Subtitle Files (.ass)", command=self.select_files)
        self.select_files_btn.pack(side=tk.LEFT, padx=5)

        self.select_dir_btn = tk.Button(button_frame, text="2. Select Output Folder", command=self.select_output_dir)
        self.select_dir_btn.pack(side=tk.LEFT, padx=5)

        self.convert_btn = tk.Button(button_frame, text="3. Convert", command=self.convert_files, state=tk.DISABLED)
        self.convert_btn.pack(side=tk.LEFT, padx=5)
        
        # Log area
        self.log_area = scrolledtext.ScrolledText(root, wrap=tk.WORD)
        self.log_area.pack(pady=10, padx=10, fill=tk.BOTH, expand=True)

    def log(self, message):
        self.log_area.insert(tk.END, message + "\n")
        self.log_area.see(tk.END)
        self.root.update_idletasks()

    def check_ready_to_convert(self):
        if self.input_files and self.output_dir:
            self.convert_btn.config(state=tk.NORMAL)
        else:
            self.convert_btn.config(state=tk.DISABLED)

    def select_files(self):
        self.input_files = filedialog.askopenfilenames(
            title="Select .ass files",
            filetypes=(("ASS files", "*.ass"), ("All files", "*.*"))
        )
        if self.input_files:
            self.log(f"Selected {len(self.input_files)} file(s):")
            for f in self.input_files:
                self.log(f"  - {os.path.basename(f)}")
        self.check_ready_to_convert()


    def select_output_dir(self):
        self.output_dir = filedialog.askdirectory(title="Select Output Folder")
        if self.output_dir:
            self.log(f"Output folder set to: {self.output_dir}")
        self.check_ready_to_convert()

    def convert_files(self):
        if not self.input_files or not self.output_dir:
            self.log("Error: Please select input files and an output directory first.")
            return

        self.log("\nStarting conversion...")
        success_count = 0
        fail_count = 0

        for file_path in self.input_files:
            base_name = os.path.basename(file_path)
            file_name_no_ext = os.path.splitext(base_name)[0]
            output_path = os.path.join(self.output_dir, f"{file_name_no_ext}.docx")
            
            self.log(f"Processing: {base_name}...")

            dialogues = parse_ass_file(file_path)
            if not dialogues:
                self.log(f"  -> Failed: No dialogue found in file.")
                fail_count += 1
                continue

            if create_docx(dialogues, output_path):
                self.log(f"  -> Success: Saved to {output_path}")
                success_count += 1
            else:
                self.log(f"  -> Failed: Could not create .docx file.")
                fail_count += 1
        
        self.log(f"\nConversion complete. {success_count} succeeded, {fail_count} failed.")
        # Reset after conversion
        self.input_files = []
        self.check_ready_to_convert()


if __name__ == "__main__":
    root = tk.Tk()
    app = SubtitleConverterApp(root)
    root.mainloop()
